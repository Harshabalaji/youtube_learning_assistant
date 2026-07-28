"""
Export service — generates PDF, DOCX, and Markdown files.
"""

import os
from pathlib import Path
from typing import Optional, Dict, Any, List

from core.config import settings
from core.logging_config import get_logger
from utils.formatters import format_complete_study_material

logger = get_logger(__name__)


def export_markdown(content: Dict[str, Any], video_title: str = "") -> str:
    """
    Export all content as a Markdown string.

    Args:
        content: Dict of content_type -> content data.
        video_title: Video title for the header.

    Returns:
        Complete markdown string.
    """
    return format_complete_study_material(content, video_title)


def export_markdown_file(
    content: Dict[str, Any],
    video_title: str = "",
    video_id: str = "",
) -> str:
    """
    Export content as a Markdown file.

    Returns:
        Path to the generated file.
    """
    export_dir = Path(settings.EXPORT_DIR)
    export_dir.mkdir(parents=True, exist_ok=True)

    filename = f"study_material_{video_id}.md"
    filepath = export_dir / filename

    markdown_content = export_markdown(content, video_title)

    with open(filepath, "w", encoding="utf-8") as f:
        f.write(markdown_content)

    logger.info("Exported Markdown: {}", filepath)
    return str(filepath)


def export_pdf(
    content: Dict[str, Any],
    video_title: str = "",
    video_id: str = "",
) -> str:
    """
    Export content as a PDF file using ReportLab.

    Returns:
        Path to the generated PDF file.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        PageBreak,
        Table,
        TableStyle,
    )
    from reportlab.lib import colors

    export_dir = Path(settings.EXPORT_DIR)
    export_dir.mkdir(parents=True, exist_ok=True)

    filename = f"study_material_{video_id}.pdf"
    filepath = export_dir / filename

    doc = SimpleDocTemplate(
        str(filepath),
        pagesize=A4,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Title"],
        fontSize=24,
        spaceAfter=20,
        textColor=HexColor("#1a1a2e"),
    )
    heading_style = ParagraphStyle(
        "CustomHeading",
        parent=styles["Heading1"],
        fontSize=16,
        spaceBefore=20,
        spaceAfter=10,
        textColor=HexColor("#16213e"),
    )
    subheading_style = ParagraphStyle(
        "CustomSubheading",
        parent=styles["Heading2"],
        fontSize=13,
        spaceBefore=12,
        spaceAfter=6,
        textColor=HexColor("#0f3460"),
    )
    body_style = ParagraphStyle(
        "CustomBody",
        parent=styles["Normal"],
        fontSize=11,
        spaceAfter=8,
        leading=16,
    )
    bullet_style = ParagraphStyle(
        "CustomBullet",
        parent=styles["Normal"],
        fontSize=11,
        spaceAfter=4,
        leftIndent=20,
        bulletIndent=10,
        leading=16,
    )

    story = []

    # Title
    story.append(Paragraph(f"📚 Study Material", title_style))
    story.append(Paragraph(video_title or "YouTube Video Analysis", heading_style))
    story.append(Spacer(1, 20))

    def add_section(title: str, text: str):
        """Helper to add a section to the PDF."""
        story.append(Paragraph(title, heading_style))
        # Split text into paragraphs and handle markdown formatting
        for line in text.split("\n"):
            line = line.strip()
            if not line:
                story.append(Spacer(1, 6))
            elif line.startswith("- ") or line.startswith("• "):
                story.append(
                    Paragraph(f"• {_escape_html(line[2:])}", bullet_style)
                )
            elif line.startswith("## "):
                story.append(Paragraph(_escape_html(line[3:]), subheading_style))
            elif line.startswith("# "):
                story.append(Paragraph(_escape_html(line[2:]), heading_style))
            else:
                story.append(Paragraph(_escape_html(line), body_style))

    # Executive Summary
    if "executive_summary" in content:
        data = content["executive_summary"]
        if isinstance(data, dict):
            add_section("📋 Executive Summary", data.get("summary", ""))

    # Key Takeaways
    if "key_takeaways" in content:
        data = content["key_takeaways"]
        if isinstance(data, dict):
            story.append(Paragraph("🎯 Key Takeaways", heading_style))
            for takeaway in data.get("takeaways", []):
                story.append(Paragraph(f"• {_escape_html(takeaway)}", bullet_style))
            story.append(Spacer(1, 10))

    # Detailed Summary
    if "detailed_summary" in content:
        data = content["detailed_summary"]
        if isinstance(data, dict):
            add_section("📝 Detailed Summary", data.get("summary", ""))

    # Notes
    if "notes" in content:
        data = content["notes"]
        if isinstance(data, dict):
            add_section("📒 Notes", data.get("markdown", ""))

    # Flashcards
    if "flashcards" in content:
        data = content["flashcards"]
        if isinstance(data, dict):
            story.append(Paragraph("🃏 Flashcards", heading_style))
            for i, card in enumerate(data.get("flashcards", []), 1):
                story.append(
                    Paragraph(
                        f"<b>Q{i}:</b> {_escape_html(card.get('question', ''))}",
                        body_style,
                    )
                )
                story.append(
                    Paragraph(
                        f"<b>A:</b> {_escape_html(card.get('answer', ''))}",
                        bullet_style,
                    )
                )
                story.append(Spacer(1, 6))

    # Vocabulary
    if "vocabulary" in content:
        data = content["vocabulary"]
        if isinstance(data, dict):
            story.append(Paragraph("📖 Vocabulary", heading_style))
            for word_entry in data.get("words", []):
                story.append(
                    Paragraph(
                        f"<b>{_escape_html(word_entry.get('word', ''))}</b>: "
                        f"{_escape_html(word_entry.get('meaning', ''))}",
                        body_style,
                    )
                )
                story.append(
                    Paragraph(
                        f"<i>Example: {_escape_html(word_entry.get('example', ''))}</i>",
                        bullet_style,
                    )
                )

    # Build PDF
    doc.build(story)
    logger.info("Exported PDF: {}", filepath)
    return str(filepath)


def export_docx(
    content: Dict[str, Any],
    video_title: str = "",
    video_id: str = "",
) -> str:
    """
    Export content as a DOCX file using python-docx.

    Returns:
        Path to the generated DOCX file.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    export_dir = Path(settings.EXPORT_DIR)
    export_dir.mkdir(parents=True, exist_ok=True)

    filename = f"study_material_{video_id}.docx"
    filepath = export_dir / filename

    doc = Document()

    # Title
    title = doc.add_heading("📚 Study Material", level=0)
    subtitle = doc.add_heading(video_title or "YouTube Video Analysis", level=1)
    doc.add_paragraph("---")

    def add_section(heading: str, text: str, level: int = 2):
        """Add a section with heading and text."""
        doc.add_heading(heading, level=level)
        for line in text.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("- ") or line.startswith("• "):
                p = doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=4)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=3)
            else:
                doc.add_paragraph(line)

    # Executive Summary
    if "executive_summary" in content:
        data = content["executive_summary"]
        if isinstance(data, dict):
            add_section("📋 Executive Summary", data.get("summary", ""))

    # Key Takeaways
    if "key_takeaways" in content:
        data = content["key_takeaways"]
        if isinstance(data, dict):
            doc.add_heading("🎯 Key Takeaways", level=2)
            for takeaway in data.get("takeaways", []):
                doc.add_paragraph(takeaway, style="List Bullet")

    # Detailed Summary
    if "detailed_summary" in content:
        data = content["detailed_summary"]
        if isinstance(data, dict):
            add_section("📝 Detailed Summary", data.get("summary", ""))

    # Notes
    if "notes" in content:
        data = content["notes"]
        if isinstance(data, dict):
            add_section("📒 Notes", data.get("markdown", ""))

    # Flashcards
    if "flashcards" in content:
        data = content["flashcards"]
        if isinstance(data, dict):
            doc.add_heading("🃏 Flashcards", level=2)
            for i, card in enumerate(data.get("flashcards", []), 1):
                p = doc.add_paragraph()
                p.add_run(f"Q{i}: ").bold = True
                p.add_run(card.get("question", ""))
                p2 = doc.add_paragraph()
                p2.add_run("A: ").bold = True
                p2.add_run(card.get("answer", ""))
                doc.add_paragraph("")

    # Interview Questions
    if "interview_questions" in content:
        data = content["interview_questions"]
        if isinstance(data, dict):
            doc.add_heading("💼 Interview Questions", level=2)
            for q in data.get("questions", []):
                p = doc.add_paragraph()
                p.add_run(f"[{q.get('difficulty', 'medium').upper()}] ").bold = True
                p.add_run(q.get("question", ""))
                p2 = doc.add_paragraph()
                p2.add_run("Suggested Answer: ").bold = True
                p2.add_run(q.get("suggested_answer", ""))
                doc.add_paragraph("")

    # Vocabulary
    if "vocabulary" in content:
        data = content["vocabulary"]
        if isinstance(data, dict):
            doc.add_heading("📖 Vocabulary", level=2)
            for word_entry in data.get("words", []):
                p = doc.add_paragraph()
                p.add_run(word_entry.get("word", "")).bold = True
                p.add_run(f": {word_entry.get('meaning', '')}")
                doc.add_paragraph(
                    f"Example: {word_entry.get('example', '')}", style="List Bullet"
                )

    doc.save(str(filepath))
    logger.info("Exported DOCX: {}", filepath)
    return str(filepath)


def _escape_html(text: str) -> str:
    """Escape HTML special characters for ReportLab."""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("**", "")
        .replace("*", "")
    )
